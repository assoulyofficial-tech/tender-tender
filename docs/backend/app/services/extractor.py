"""
Text Extraction Pipeline.
Extracts text from PDF, DOC/DOCX, XLS/XLSX documents.
Uses OCR (PaddleOCR) only when text extraction fails.
No AI calls - pure text extraction.
"""

import io
import logging
from dataclasses import dataclass
from typing import Optional
from enum import Enum

logger = logging.getLogger(__name__)


class ExtractionMethod(str, Enum):
    DIGITAL = "digital"  # Direct text extraction
    OCR = "ocr"  # PaddleOCR fallback


@dataclass
class ExtractionResult:
    """Result of text extraction."""
    text: str
    method: ExtractionMethod
    page_count: Optional[int] = None
    success: bool = True
    error: Optional[str] = None


class TextExtractor:
    """
    Document text extractor with OCR fallback.
    
    Supported formats:
    - PDF (digital and scanned)
    - DOC/DOCX
    - XLS/XLSX
    
    Uses PaddleOCR locally (CPU) when digital extraction fails.
    """
    
    def __init__(self):
        self._ocr = None  # Lazy load PaddleOCR
    
    @property
    def ocr(self):
        """Lazy load PaddleOCR for CPU."""
        if self._ocr is None:
            from paddleocr import PaddleOCR
            self._ocr = PaddleOCR(
                use_angle_cls=True,
                lang='fr',  # French language
                use_gpu=False,  # CPU only
                show_log=False,
            )
        return self._ocr
    
    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        """
        Extract text from document.
        
        Args:
            content: Document bytes
            filename: Original filename (for format detection)
            
        Returns:
            ExtractionResult with text and metadata
        """
        ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
        
        try:
            if ext == 'pdf':
                return self._extract_pdf(content)
            elif ext in ('doc', 'docx'):
                return self._extract_docx(content, ext)
            elif ext in ('xls', 'xlsx'):
                return self._extract_excel(content, ext)
            else:
                return ExtractionResult(
                    text='',
                    method=ExtractionMethod.DIGITAL,
                    success=False,
                    error=f"Unsupported format: {ext}"
                )
        except Exception as e:
            logger.error(f"Extraction failed for {filename}: {e}")
            return ExtractionResult(
                text='',
                method=ExtractionMethod.DIGITAL,
                success=False,
                error=str(e)
            )
    
    def _extract_pdf(self, content: bytes) -> ExtractionResult:
        """Extract text from PDF, with OCR fallback."""
        import fitz  # PyMuPDF
        
        doc = fitz.open(stream=content, filetype="pdf")
        page_count = len(doc)
        
        # First attempt: digital extraction
        text_parts = []
        has_text = False
        
        for page in doc:
            page_text = page.get_text("text")
            text_parts.append(page_text)
            if page_text.strip():
                has_text = True
        
        doc.close()
        
        # If we got meaningful text, return it
        full_text = "\n".join(text_parts).strip()
        if has_text and len(full_text) > 50:  # Minimum threshold
            return ExtractionResult(
                text=full_text,
                method=ExtractionMethod.DIGITAL,
                page_count=page_count,
                success=True
            )
        
        # Fallback to OCR for scanned PDFs
        logger.info(f"PDF has no/little text, using OCR (pages: {page_count})")
        return self._ocr_pdf(content, page_count)
    
    def _ocr_pdf(self, content: bytes, page_count: int) -> ExtractionResult:
        """OCR a PDF using PaddleOCR."""
        import fitz
        
        doc = fitz.open(stream=content, filetype="pdf")
        all_text = []
        
        for page_num, page in enumerate(doc):
            # Convert page to image
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better OCR
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            
            # OCR the image
            page_text = self._ocr_image(img_bytes)
            if page_text:
                all_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        doc.close()
        
        return ExtractionResult(
            text="\n\n".join(all_text),
            method=ExtractionMethod.OCR,
            page_count=page_count,
            success=True
        )
    
    def _ocr_image(self, img_bytes: bytes) -> str:
        """OCR a single image using PaddleOCR."""
        import numpy as np
        from PIL import Image
        
        # Convert bytes to numpy array
        img = Image.open(io.BytesIO(img_bytes))
        img_array = np.array(img)
        
        # Run OCR
        result = self.ocr.ocr(img_array, cls=True)
        
        if not result or not result[0]:
            return ""
        
        # Extract text from result
        lines = []
        for line in result[0]:
            if line and len(line) >= 2:
                text = line[1][0]  # text content
                lines.append(text)
        
        return "\n".join(lines)
    
    def _extract_docx(self, content: bytes, ext: str) -> ExtractionResult:
        """Extract text from DOC/DOCX."""
        if ext == 'docx':
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)
            
            text = "\n".join(paragraphs)
            
            return ExtractionResult(
                text=text,
                method=ExtractionMethod.DIGITAL,
                success=True
            )
        else:
            # DOC format - try antiword or similar
            # For now, return error suggesting conversion
            return ExtractionResult(
                text='',
                method=ExtractionMethod.DIGITAL,
                success=False,
                error="DOC format not supported. Please convert to DOCX."
            )
    
    def _extract_excel(self, content: bytes, ext: str) -> ExtractionResult:
        """Extract text from XLS/XLSX."""
        import openpyxl
        
        if ext == 'xlsx':
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        else:
            # XLS format - use xlrd
            import xlrd
            xls_book = xlrd.open_workbook(file_contents=content)
            
            text_parts = []
            for sheet_idx in range(xls_book.nsheets):
                sheet = xls_book.sheet_by_index(sheet_idx)
                text_parts.append(f"=== Sheet: {sheet.name} ===")
                
                for row_idx in range(sheet.nrows):
                    row_values = []
                    for col_idx in range(sheet.ncols):
                        cell_value = sheet.cell_value(row_idx, col_idx)
                        if cell_value:
                            row_values.append(str(cell_value))
                    if row_values:
                        text_parts.append("\t".join(row_values))
            
            return ExtractionResult(
                text="\n".join(text_parts),
                method=ExtractionMethod.DIGITAL,
                success=True
            )
        
        # XLSX extraction
        text_parts = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    text_parts.append("\t".join(row_values))
        
        wb.close()
        
        return ExtractionResult(
            text="\n".join(text_parts),
            method=ExtractionMethod.DIGITAL,
            success=True
        )


# Global extractor instance
text_extractor = TextExtractor()
